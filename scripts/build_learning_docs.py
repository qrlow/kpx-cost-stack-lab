from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_FUNDAMENTALS = "course/Korea_Power_Stack_Primer.docx"
OUT_WALKTHROUGH = "course/Model_Walkthrough.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(36, 35, 31)
MUTED = RGBColor(92, 92, 92)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"

SOURCES = [
    ("S1", "Korea Power Exchange (KPX), 2025 Power Market Statistics, published 2026-03-31.",
     "https://www.kpx.or.kr/boardDownload.es?bid=0045&list_no=77084&seq=4"),
    ("S2", "KPX, Electricity Market Trading Process.",
     "https://www.kpx.or.kr/menu.es?mid=a20201000000"),
    ("S3", "KPX, Market Price Determination.",
     "https://www.kpx.or.kr/menu.es?mid=a20203000000"),
    ("S4", "KPX, Electricity Market Trading System.",
     "https://www.kpx.or.kr/menu.es?mid=a20202000000"),
    ("S5", "KPX, Electricity Market Operation Council / Cost Evaluation Committee.",
     "https://new.kpx.or.kr/menu.es?mid=a20205000000"),
    ("S6", "KPX, Settlement Process.",
     "https://new.kpx.or.kr/menu.es?mid=a20204000000"),
    ("S7", "IEA, Korea Electricity Security Review.",
     "https://www.iea.org/reports/korea-electricity-security-review"),
    ("S8", "MOTIE, 11th Basic Plan for Long-term Electricity Supply and Demand announcement.",
     "https://www.motir.go.kr/kor/article/ATCLc01b2801b/70083/view"),
    ("S9", "KDI Economic Information and Education Center, MOTIE 11th Basic Plan confirmation release.",
     "https://eiec.kdi.re.kr/policy/materialView.do?num=263534"),
    ("S10", "Ministry of Climate, Energy and Environment, Fine Dust Seasonal Management release.",
     "https://mcee.go.kr/eng/web/board/read.do?boardId=1640870&boardMasterId=522&menuId=461"),
    ("S11", "ICAP, Korea Emissions Trading System factsheet.",
     "https://icapcarbonaction.com/en/ets/korea-emissions-trading-system-k-ets"),
    ("S12", "IEA, ETS in the power sector, Korea case discussion.",
     "https://www.iea.org/reports/implementing-effective-emissions-trading-systems/ets-in-power-sector"),
    ("S13", "IEA, Natural Gas Supply Security in Korea.",
     "https://iea.blob.core.windows.net/assets/4240647e-42f1-4791-8b12-59a09ca8fff9/NaturalGasSupplySecurityinKorea.pdf"),
    ("S14", "IEA, Korea 2025 Energy Policy Review, executive summary.",
     "https://www.iea.org/reports/korea-2025/executive-summary"),
    ("S15", "MarketWatch / Dow Jones OPIS, South Korea Approves Market Stability Reserve for K-ETS.",
     "https://www.marketwatch.com/story/south-korea-approves-market-stability-reserve-for-k-ets-opis-31412180"),
    ("S16", "GitHub repository for Monthly Burn.",
     "https://github.com/qrlow/kpx-cost-stack-lab"),
]


def set_run(run, size=None, bold=None, italic=None, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))


def configure_doc(doc, label):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    hp = section.header.paragraphs[0]
    hp.text = label
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(hp.runs[0], size=9, color=MUTED)
    fp = section.footer.paragraphs[0]
    fp.text = "Prepared for Korea thermal coal market analysis | 2026-06-25"
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(fp.runs[0], size=9, color=MUTED)


def add_para(doc, text="", style=None, bold_first=None):
    p = doc.add_paragraph(style=style)
    if bold_first and text.startswith(bold_first):
        r = p.add_run(bold_first)
        set_run(r, bold=True)
        r2 = p.add_run(text[len(bold_first):])
        set_run(r2)
    else:
        r = p.add_run(text)
        set_run(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run(r)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run(r)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run(r, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run(r2)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run(r, font="Consolas", size=10, color=RGBColor(40, 40, 40))


def add_table(doc, headers, rows, widths, source=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    mark_header_row(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run(r, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run(r)
    if source:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("Source: " + source)
        set_run(r, size=9, color=MUTED, italic=True)
    return table


def add_cover(doc, title, subtitle, metadata):
    for _ in range(3):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = kicker.add_run("Korean Thermal Coal Market")
    set_run(r, size=12, bold=True, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run(r, size=26, bold=True, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(8)
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sp.add_run(subtitle)
    set_run(r, size=13, color=MUTED)
    doc.add_paragraph()
    add_table(doc, ["Field", "Value"], metadata, [1.5, 5.0])
    doc.add_page_break()


def add_references(doc, source_ids):
    doc.add_heading("References", level=1)
    source_map = {sid: (title, url) for sid, title, url in SOURCES}
    for sid in source_ids:
        title, url = source_map[sid]
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"[{sid}] {title} ")
        set_run(r, bold=True)
        r2 = p.add_run(url)
        set_run(r2, color=RGBColor(0, 102, 204))


def build_fundamentals():
    doc = Document()
    configure_doc(doc, "Korea Power Stack Primer")
    add_cover(doc, "Korea Power Stack Primer",
              "A compact primer on Korea's power-market structure, coal/LNG switching drivers, and why KPX dispatch matters for thermal coal.",
              [
                  ("Audience", "Thermal coal trader or analyst"),
                  ("Purpose", "Understand Korea power-market fundamentals before reading scenario outputs"),
                  ("Scope", "Institutions, KPX cost-based pool, 2025 fuel data, policy overlays, and trader implications"),
                  ("Date", "2026-06-25"),
              ])

    doc.add_heading("Executive Read", level=1)
    add_callout(doc, "Core market point",
                "Korea is not a simple merchant coal-versus-JKM switch. KPX runs a centrally scheduled Cost-Based Pool; KEPCO is the single buyer; generator costs are assessed; and coal burn is shaped by nuclear output, LNG availability, fine-dust restrictions, carbon recognition, and system constraints as much as by spot coal price.")
    add_bullets(doc, [
        "The commercial coal question is not just whether coal is cheaper than LNG. The better question is whether coal can physically and legally run, and whether its recognized cost position is changing versus LNG.",
        "Coal produced 164.8 TWh of KPX-traded electricity in 2025, close to LNG at 159.2 TWh. That makes small monthly shifts meaningful for prompt cargo demand. [S1]",
        "LNG set the mainland SMP in most counted 2025 intervals, but coal still carried large baseload energy. Price setting and fuel burn are different signals. [S1]",
        "Winter fine-dust policy can restrict coal even when coal is in merit; nuclear outages can lift fossil demand even when coal/LNG spreads do not move. [S10], [S14]",
    ])

    doc.add_heading("1. Market Architecture", level=1)
    add_para(doc, "KPX operates Korea's wholesale power market and system operation. KPX describes the market as a pool in which generators sell electricity and KEPCO acts as the single buyer. The IEA describes the structure as a day-ahead wholesale market run by KPX, with KEPCO controlling transmission, distribution, and retail supply. [S2], [S7]")
    add_table(doc, ["Actor", "What they control", "Why a coal trader should care"], [
        ("KPX", "Market operation, system scheduling, SMP calculation, settlement administration", "Determines day-ahead schedules and marginal pricing under cost-based rules."),
        ("KEPCO", "Single buyer, transmission/distribution/retail monopoly role", "Retail pricing and wholesale cost recovery do not work like a fully liberalized merchant market."),
        ("KEPCO gencos", "Large coal and LNG fleet ownership", "Their outages, fuel costs, and policy compliance affect aggregate coal burn."),
        ("KHNP", "Nuclear fleet", "Nuclear output is a major residual-load driver for fossil generation."),
        ("KOGAS/direct LNG importers", "Gas procurement and supply channel", "LNG economics may reflect portfolio procurement or direct-import economics, not spot JKM alone."),
        ("MOTIE/MCEE", "Long-term supply plan, energy policy, fine-dust policy", "Can shift coal availability independently of price spreads."),
    ], [1.25, 2.15, 3.10], "KPX and IEA market structure descriptions. [S2], [S7]")

    doc.add_heading("2. KPX Cost-Based Pool", level=1)
    add_para(doc, "Korea's wholesale market is a Cost-Based Pool. KPX states that generator fixed costs and variable costs are examined monthly by the Generation Cost Assessment Committee based on documents submitted by generators. The Cost Evaluation Committee performs the assessed-cost function. [S2], [S5]")
    add_para(doc, "KPX determines market prices one day ahead for each hour. Its Price Setting Schedule selects generator operation to meet forecast demand while satisfying system constraints at the lowest generation cost; the highest-cost dispatched generator becomes the marginal plant and sets SMP. [S3]")
    add_para(doc, "The trading-system description matters because KPX explicitly refers to schedules that consider fuel and transmission constraints. A pure coal/Japan-Korea Marker comparison misses that operational layer. [S4]")
    add_table(doc, ["Term", "Meaning", "Trader relevance"], [
        ("GCAC", "Monthly assessment of generator costs", "Fuel-price changes can lag in recognized dispatch economics."),
        ("PSS", "Price Setting Schedule", "The relevant stack is assessed-cost dispatch, not free bids."),
        ("SMP", "System Marginal Price", "Frequently set by LNG, but not equal to LNG generation share."),
        ("Settlement", "Post-trade cash process", "Settlement has more components than energy price; burn and payments are not the same thing."),
    ], [1.15, 2.55, 2.8], "KPX trading, price-determination, and settlement process pages. [S2]-[S6]")

    doc.add_heading("3. 2025 Fuel Baseline", level=1)
    add_table(doc, ["Fuel", "Capacity MW", "Trading GWh", "Share of trading volume"], [
        ("Nuclear", "26,050", "175,549", "32.1%"),
        ("Coal", "41,739", "164,790", "30.1%"),
        ("LNG", "48,388", "159,227", "29.1%"),
        ("Renewables", "19,457", "39,697", "7.3%"),
        ("Pumped storage", "4,700", "4,414", "0.8%"),
        ("Oil / other", "5,824", "3,485", "0.6%"),
    ], [1.35, 1.45, 1.45, 2.25], "KPX 2025 Power Market Statistics. [S1]")
    add_para(doc, "The key trading read is that coal and LNG were close in annual traded energy. Coal was not marginal most of the time, but it was still a large physical baseload consumer. That distinction is where a coal trader can have edge: SMP tells you marginal price setting; monthly fuel-burn data tells you physical demand.")
    add_table(doc, ["Month", "Coal GWh", "LNG GWh", "What it suggests"], [
        ("January", "12,709", "15,345", "Winter coal restrictions and LNG/heating season matter."),
        ("April", "7,936", "13,198", "Shoulder-month coal burn was low versus summer."),
        ("July", "19,582", "13,720", "Summer demand/nuclear profile supported coal burn."),
        ("August", "20,208", "14,853", "Highest 2025 monthly coal burn in the KPX baseline."),
        ("December", "16,660", "13,261", "Winter demand supported coal, but seasonal policy can still cap upside."),
    ], [1.0, 1.1, 1.1, 3.3], "KPX monthly electric energy trading volume by fuel type. [S1]")

    doc.add_heading("4. What Moves Korean Coal Burn", level=1)
    add_table(doc, ["Driver", "Bullish for coal burn when...", "Bearish / switch risk when..."], [
        ("Power demand", "Weather or industrial load lifts monthly fossil need.", "Demand softens in shoulder months."),
        ("Nuclear output", "Outages reduce low-cost baseload and raise residual fossil demand.", "High nuclear availability compresses coal and LNG burn."),
        ("Renewables", "Weak solar/wind increases residual demand.", "Strong renewables reduce fossil need, especially around solar-heavy hours."),
        ("Coal restrictions", "Caps relax or seasonal controls are not binding.", "Fine-dust season constrains coal units regardless of fuel spread."),
        ("LNG price", "LNG rises relative to coal, widening coal's merit advantage.", "LNG falls enough to displace the switchable coal band."),
        ("FX", "KRW strength lowers imported-fuel cost; relative effect depends on commodity move.", "KRW weakness raises both imported coal and LNG costs."),
        ("Carbon recognition", "Carbon is low or not fully reflected in dispatch economics.", "Carbon is recognized strongly; coal is more CO2-intensive than LNG."),
    ], [1.2, 2.55, 2.75])

    doc.add_heading("5. Fuel Economics Translation", level=1)
    add_para(doc, "Coal traders think in USD/t and calorific value, while KPX-style dispatch economics require KRW/kWh. The bridge is energy content, heat rate, FX, variable O&M, and carbon.")
    add_code(doc, "coal USD/MMBtu = coal USD/t / MMBtu per tonne")
    add_code(doc, "fuel KRW/kWh = USD/MMBtu * USD/KRW * heat rate MMBtu/MWh / 1000")
    add_code(doc, "carbon KRW/kWh = tCO2e/MWh * KRW/tCO2e * carbon recognition % / 100 / 1000")
    add_para(doc, "MOTIE's 11th Basic Plan materials provide Korea-specific average power-sector emissions factors used for coal and LNG comparison: coal 0.8384 tCO2e/MWh and LNG 0.3800 tCO2e/MWh. [S8]")

    doc.add_heading("6. Policy and Structural Overlays", level=1)
    add_para(doc, "Fine-dust seasonal management is important because it can make coal availability a policy variable. The English ministry release describes stronger measures from December 1 to March 31, including coal-unit shutdowns and output caps in the referenced plan. [S10]")
    add_para(doc, "K-ETS is also relevant, but carbon cost does not always translate one-for-one into dispatch in a regulated cost-based market. ICAP describes K-ETS coverage and Phase 4 rules; the IEA highlights that Korea's cost-based market historically limited direct carbon-price reflection in dispatch and retail signals. [S11], [S12]")
    add_para(doc, "Long-term supply planning matters for the trade because more nuclear and renewables structurally reduce fossil residual demand, while coal retirements/conversions reduce coal's upper bound. The 11th Basic Plan covers 2024-2038 and includes new nuclear capacity, renewable expansion, and a shifting generation mix. [S8], [S9]")

    doc.add_heading("7. Questions a Trader Will Ask", level=1)
    add_bullets(doc, [
        "Is this burn or marginal price? Answer: treat KPX monthly fuel generation as physical burn proxy; SMP tells marginal price-setting, not coal consumption.",
        "Does this use customs imports? Answer: no. It uses KPX-traded generation; imports/inventory/procurement timing are separate layers.",
        "Where can edge come from? Answer: nuclear outage expectations, fine-dust/cap assumptions, LNG relative value, and month-specific fossil residual demand.",
        "Why not just use spark/dark spread? Answer: Korea's cost-based dispatch, recognized-cost timing, coal floors/caps, and policy restrictions make a simple spread incomplete.",
    ])
    add_references(doc, ["S1", "S2", "S3", "S4", "S5", "S6", "S7", "S8", "S9", "S10", "S11", "S12", "S13", "S14"])
    doc.core_properties.title = "Korea Power Stack Primer"
    doc.core_properties.subject = "Korean power-market primer for thermal coal analysis"
    doc.core_properties.author = "OpenAI Codex"
    doc.save(OUT_FUNDAMENTALS)


def build_walkthrough():
    doc = Document()
    configure_doc(doc, "Model Walkthrough")
    add_cover(doc, "Model Walkthrough",
              "How Monthly Burn works, with trader-focused definitions and interpretation guidance.",
              [
                  ("Audience", "Thermal coal trader or analyst"),
                  ("Tools", "Monthly Burn"),
                  ("App URL", "https://qrlow.github.io/kpx-cost-stack-lab/"),
                  ("Date", "2026-06-25"),
              ])

    doc.add_heading("One-Minute Pitch", level=1)
    add_callout(doc, "What the project is really showing",
                "Monthly Burn converts Korean power-market fundamentals into a coal trader's language: tonnes, cargoes, 3-month demand, LNG switch risk, and coal price headroom. It is not trying to be a black-box forecast. It is a transparent scenario engine that shows which market lever would change Korean thermal coal demand.")
    add_bullets(doc, [
        "Default settings reproduce the selected 2025 KPX monthly coal/LNG generation baseline instead of forcing all coal to run just because it is cheaper on paper.",
        "Scenario changes show incremental coal burn versus that baseline, which is the part a trader can map to tender timing, cargo demand, and relative-value risk.",
        "The model separates physical constraints from price incentives: coal can be cheap and still capped; LNG can be expensive and still needed for residual load.",
    ])

    doc.add_heading("1. Monthly Model: What It Answers", level=1)
    add_para(doc, "The monthly model asks: for a selected delivery month, how much coal-fired generation does Korea need versus its 2025 KPX baseline, and what does that imply for thermal coal tonnes and cargo equivalents?")
    add_table(doc, ["Output", "What it means", "Trader use"], [
        ("Coal burn", "Modeled coal-fired generation in GWh and converted million tonnes", "Physical demand read, not just price signal."),
        ("Cargo equivalent", "Coal tonnes divided by selected cargo size", "Turns burn delta into prompt cargo/tender scale."),
        ("3-month strip", "Selected month plus next two months", "Matches procurement windows better than a single-month view."),
        ("Coal headroom", "Break-even coal price versus LNG minus current coal price", "Measures how much coal can rally before LNG starts threatening the switchable band."),
        ("Fuel cost spread", "LNG variable cost minus coal variable cost", "Positive means coal is cheaper; negative means LNG is cheaper."),
        ("Trader signal", "Balanced, tender support, strong coal bid, soft demand, or LNG switch risk", "Quick screening label for the scenario."),
    ], [1.25, 2.45, 2.8])

    doc.add_heading("2. Monthly Engine Step-by-Step", level=1)
    add_numbers(doc, [
        "Start from actual 2025 KPX monthly traded generation by fuel: nuclear, coal, LNG, renewables, pumped storage, oil, other, and total load proxy. [S1]",
        "Apply scenario shocks to total demand, nuclear output, and renewable output. These determine monthly fossil residual need.",
        "Keep oil, pumped storage, and other generation as fixed baseline blocks for simplicity.",
        "Calculate fossil need: total power demand minus nuclear, renewables, and fixed other generation.",
        "Calculate coal maximum energy from coal capacity, month hours, coal availability, and any Dec-Mar seasonal restriction stress.",
        "Calculate coal floor energy from baseline coal burn times the Coal Floor percentage. This is the physical-inflexibility anchor.",
        "Allocate incremental fossil need to coal or LNG based on variable cost order, subject to coal max, coal floor, and LNG availability.",
        "Apply economic switching only to the switchable fossil band, and only when coal's relative advantage changes versus the default reference case.",
        "Convert coal GWh to tonnes and cargoes, then aggregate the selected month and next two months for the strip view.",
    ])

    doc.add_heading("3. Physical Inflexibility of Coal Units", level=1)
    add_para(doc, "Coal plants cannot behave like perfectly flexible financial spread options. They have minimum stable generation, slower start/stop dynamics, planned maintenance, fuel scheduling, environmental limits, and operating commitments. The monthly model captures this with two constraints.")
    add_table(doc, ["Constraint", "Model representation", "Interpretation"], [
        ("Coal maximum", "coal capacity MW * 24 * days * coal availability %; optionally derated in Dec-Mar when seasonal restriction stress is on", "Coal cannot exceed available fleet energy even if it is very cheap."),
        ("Coal floor", "baseline monthly coal GWh * Coal Floor %", "A portion of coal burn is treated as sticky and not easily displaced by LNG."),
        ("Switchable fossil band", "(baseline coal + baseline LNG) * Switchable Fossil Band %", "Only this band can move on relative fuel economics in the monthly screen."),
        ("Switch intensity", "absolute shift in coal advantage / 12 KRW/kWh, capped at 100%", "Small price changes move part of the band; large price changes move all contestable energy."),
    ], [1.35, 2.55, 2.6])
    add_callout(doc, "How to explain this to a trader",
                "The model does not say every coal MWh switches the moment LNG gets cheaper. It says there is a sticky coal base, a physical coal ceiling, and a contestable fossil band. That is closer to how monthly burn and procurement actually behave.")

    doc.add_heading("4. Fuel Cost Spread", level=1)
    add_para(doc, "Fuel cost spread is the model's simple coal-versus-LNG merit read.")
    add_code(doc, "fuel cost spread = LNG variable cost KRW/kWh - coal variable cost KRW/kWh")
    add_bullets(doc, [
        "Positive spread: LNG is more expensive than coal, so coal has dispatch-cost advantage.",
        "Negative spread: LNG is cheaper than coal, so LNG switch risk rises.",
        "It includes fuel, heat rate, variable O&M, FX, carbon price, and carbon recognition.",
        "It is not the same thing as observed SMP. SMP is a market price; this spread is a relative variable-cost screen.",
    ])

    doc.add_heading("5. Coal Headroom", level=1)
    add_para(doc, "Coal headroom is the most trader-friendly price output. It answers: how far can the coal price rise before modeled coal loses its variable-cost advantage to LNG?")
    add_code(doc, "coal headroom USD/t = break-even coal USD/t versus LNG - current coal USD/t")
    add_bullets(doc, [
        "Positive headroom means coal is still cheaper than LNG by that many USD/t under the scenario.",
        "Near-zero headroom means the trade is close to the LNG switch boundary.",
        "Negative headroom means coal is out of merit versus LNG before considering physical floors, caps, or policy constraints.",
        "Headroom changes with LNG price, FX, coal NAR, heat rates, carbon price, and carbon recognition.",
    ])
    add_para(doc, "The break-even calculation backs out the coal fuel price that would make coal variable cost equal LNG variable cost after coal O&M and recognized carbon are included.")

    doc.add_heading("6. Cargo Equivalent", level=1)
    add_para(doc, "Cargo equivalent converts power burn into a unit coal traders actually think about.")
    add_code(doc, "coal Mt = coal GWh * 1000 MWh/GWh * coal heat rate / MMBtu per tonne / 1,000,000")
    add_code(doc, "cargo equivalents = coal Mt * 1000 / cargo size kt")
    add_para(doc, "If the model shows +0.45 Mt versus baseline and the cargo size is 150 kt, that is roughly +3 cargoes. This is a scale indicator, not a named tender forecast. It helps decide whether a scenario is commercially material.")

    doc.add_heading("7. 3-Month Strip", level=1)
    add_para(doc, "The 3-month strip is the selected month plus the next two months. If August is selected, the strip is August/September/October. This approximates how traders think about procurement and delivery windows better than a single calendar month.")
    add_bullets(doc, [
        "Use the monthly number for immediate burn sensitivity.",
        "Use the 3-month strip for cargo program pressure and whether the story is persistent.",
        "A one-month spike may be weather/noise; a strip increase is more relevant to procurement and relative value.",
    ])

    doc.add_heading("8. Monthly Controls", level=1)
    add_table(doc, ["Control", "What it changes", "Question it answers"], [
        ("Delivery Month", "Selects KPX 2025 monthly baseline and days in month", "Is this a summer, winter, or shoulder-month setup?"),
        ("Power Demand vs 2025", "Scales total monthly demand", "Is weather/industrial load adding fossil burn?"),
        ("Nuclear Output vs 2025", "Scales nuclear generation", "Are outages lifting coal/LNG residual demand?"),
        ("Renewables vs 2025", "Scales renewable generation", "Is weak renewable output adding fossil demand?"),
        ("Coal Availability", "Caps maximum coal energy", "Can coal physically run more?"),
        ("Coal Floor", "Defines sticky coal burn", "How much coal is hard to displace?"),
        ("Switchable Fossil Band", "Defines coal/LNG volume exposed to price switching", "How much burn is actually contestable?"),
        ("Coal/LNG/FX/Carbon", "Sets variable-cost comparison", "Is coal still in merit versus LNG?"),
        ("Cargo Size", "Converts Mt into cargo count", "What is the tender/cargo scale of the change?"),
        ("Seasonal Restriction Stress", "Applies Dec-Mar coal derate", "Does fine-dust season cap coal upside?"),
    ], [1.45, 2.45, 2.6])

    doc.add_heading("9. Model Interpretation Notes", level=1)
    add_table(doc, ["Topic", "Interpretation"], [
        ("Is this a forecast?", "No. It is a transparent scenario model anchored to KPX 2025 monthly data. Its value is decomposition: what lever changes Korean coal burn and by how much."),
        ("Why not customs import data?", "Imports reflect procurement timing, inventory, quality, and stock changes. KPX burn is closer to power-sector consumption; imports are the next layer to add."),
        ("Why does default match 2025 instead of maxing cheap coal?", "Because the baseline already embeds real constraints and operations. Only changes versus the reference case move the contestable band."),
        ("What does negative headroom mean?", "At the chosen LNG, FX, carbon, and heat-rate assumptions, coal is above the LNG break-even and exposed to LNG displacement, subject to coal floor and constraints."),
        ("What is the biggest weakness?", "No unit-level outage stack, no inventory/tender data, and no live price feeds yet. Those are obvious upgrades rather than hidden assumptions."),
        ("Why is this useful?", "It turns a Korea power-market story into trader units: Mt, cargoes, strip demand, and switch risk."),
    ], [1.65, 4.85])

    doc.add_heading("10. Current Model Assumptions", level=1)
    add_table(doc, ["Assumption", "Current value", "Rationale"], [
        ("Coal heat rate", "9.15 MMBtu/MWh", "Representative monthly fleet heat rate used for burn conversion."),
        ("LNG heat rate", "7.10 MMBtu/MWh", "Representative combined-cycle monthly comparison point."),
        ("Coal emissions factor", "0.8384 tCO2e/MWh", "MOTIE Korea-specific average factor. [S8]"),
        ("LNG emissions factor", "0.3800 tCO2e/MWh", "MOTIE Korea-specific average factor. [S8]"),
        ("Default carbon price", "KRW 16,700/tCO2e", "Illustrative OPIS/MarketWatch KAU25 reference. [S15]"),
        ("Contestable switch threshold", "12 KRW/kWh for full switch band", "Screening parameter; not a KPX rule."),
        ("Procurement window", "3 months", "Trader-oriented strip view."),
    ], [1.8, 1.5, 3.2])

    doc.add_heading("11. Limitations and Upgrade Path", level=1)
    add_bullets(doc, [
        "Not a confidential KPX production-cost model.",
        "No unit-level outage schedule, ramp-rate optimization, minimum up/down time, or network constraints.",
        "No customs imports, utility inventories, or tender-by-tender procurement layer.",
        "No automated live coal/LNG/FX/carbon data feed yet.",
        "Best next upgrade: add import/inventory/tender data and a live commodity-price snapshot, then compare modeled burn pressure to actual procurement behavior.",
    ])
    add_references(doc, ["S1", "S2", "S3", "S4", "S8", "S10", "S11", "S12", "S15", "S16"])
    doc.core_properties.title = "Model Walkthrough"
    doc.core_properties.subject = "Monthly Burn model methodology"
    doc.core_properties.author = "OpenAI Codex"
    doc.save(OUT_WALKTHROUGH)


if __name__ == "__main__":
    build_fundamentals()
    build_walkthrough()
