from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "course/KPX_Cost_Stack_Lab_Course.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(36, 35, 31)
MUTED = RGBColor(92, 92, 92)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
WHITE = "FFFFFF"


SOURCES = [
    ("S1", "Korea Power Exchange (KPX), 2025 Power Market Statistics, published 2026-03-31.",
     "https://www.kpx.or.kr/boardDownload.es?bid=0045&list_no=77084&seq=4"),
    ("S2", "KPX, Electricity Market Trading Process.",
     "https://www.kpx.or.kr/menu.es?mid=a20201000000"),
    ("S3", "KPX, Market Price Determination.",
     "https://www.kpx.or.kr/menu.es?mid=a20203000000"),
    ("S4", "KPX, Electricity Market Trading System.",
     "https://www.kpx.or.kr/menu.es?mid=a20202000000"),
    ("S5", "KPX, Electricity Market Operation Council / Cost Evaluation Committee.",
     "https://new.kpx.or.kr/menu.es?mid=a20205000000"),
    ("S6", "KPX, Settlement Process.",
     "https://new.kpx.or.kr/menu.es?mid=a20204000000"),
    ("S7", "International Energy Agency (IEA), Korea Electricity Security Review.",
     "https://www.iea.org/reports/korea-electricity-security-review"),
    ("S8", "MOTIE, 11th Basic Plan for Long-term Electricity Supply and Demand (2024-2038) announcement.",
     "https://www.motir.go.kr/kor/article/ATCLc01b2801b/70083/view"),
    ("S9", "KDI Economic Information and Education Center, MOTIE 11th Basic Plan confirmation release.",
     "https://eiec.kdi.re.kr/policy/materialView.do?num=263534"),
    ("S10", "Ministry of Climate, Energy and Environment, Fine Dust Seasonal Management press release.",
     "https://mcee.go.kr/eng/web/board/read.do?boardId=1640870&boardMasterId=522&menuId=461"),
    ("S11", "ICAP, Korea Emissions Trading System factsheet.",
     "https://icapcarbonaction.com/en/ets/korea-emissions-trading-system-k-ets"),
    ("S12", "IEA, ETS in the power sector, Korea case discussion.",
     "https://www.iea.org/reports/implementing-effective-emissions-trading-systems/ets-in-power-sector"),
    ("S13", "IEA, Natural Gas Supply Security in Korea.",
     "https://iea.blob.core.windows.net/assets/4240647e-42f1-4791-8b12-59a09ca8fff9/NaturalGasSupplySecurityinKorea.pdf"),
    ("S14", "IEA, Korea 2025 Energy Policy Review, executive summary.",
     "https://www.iea.org/reports/korea-2025/executive-summary"),
    ("S15", "IEA, Coal 2024.",
     "https://iea.blob.core.windows.net/assets/a1ee7b75-d555-49b6-b580-17d64ccc8365/Coal2024.pdf"),
    ("S16", "MarketWatch / Dow Jones OPIS, South Korea Approves Market Stability Reserve for K-ETS.",
     "https://www.marketwatch.com/story/south-korea-approves-market-stability-reserve-for-k-ets-opis-31412180"),
    ("S17", "KPX mainland SMP live table.",
     "https://www.kpx.or.kr/smpInland.es?mid=a10404080100&device=pc"),
    ("S18", "KPX daily power supply-demand performance page.",
     "https://www.kpx.or.kr/powerDemandPerform.es?mid=a10404060000"),
    ("S19", "GitHub repository for the KPX Cost Stack Lab model.",
     "https://github.com/qrlow/kpx-cost-stack-lab"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_w = row.cells[idx]._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                row.cells[idx]._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_run(run, size=None, bold=None, italic=None, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_para(doc, text="", style=None, bold_first=None, keep=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.keep_together = keep
    if bold_first and text.startswith(bold_first):
        r = p.add_run(bold_first)
        set_run(r, bold=True)
        r2 = p.add_run(text[len(bold_first):])
        set_run(r2)
    else:
        r = p.add_run(text)
        set_run(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run(r)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run(r)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run(r, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run(r2)


def add_table(doc, headers, rows, widths, source=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    mark_header_row(table.rows[0])
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(hdr[i], LIGHT_BLUE)
        set_cell_margins(hdr[i])
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run(r)
    if source:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("Source: " + source)
        set_run(r, size=9, color=MUTED, italic=True)
    return table


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run(r, font="Consolas", size=10, color=RGBColor(40, 40, 40))


def set_header_footer(doc):
    for section in doc.sections:
        hp = section.header.paragraphs[0]
        hp.text = "KPX Cost Stack Lab Course"
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.runs[0].font.size = Pt(9)
        hp.runs[0].font.color.rgb = MUTED
        fp = section.footer.paragraphs[0]
        fp.text = "Generated for Korean thermal coal market learning | 2026-06-25"
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.runs[0].font.size = Pt(9)
        fp.runs[0].font.color.rgb = MUTED


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = kicker.add_run("Technical Course and Model Guide")
    set_run(r, size=12, bold=True, color=MUTED)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Understanding Korea's KPX Cost Stack and Thermal Coal-to-LNG Switching")
    set_run(r, size=26, bold=True, color=DARK_BLUE)
    title.paragraph_format.space_after = Pt(8)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("A university-level guide to the market design, data, equations, assumptions, and policy overlays behind the KPX Cost Stack Lab")
    set_run(r, size=13, color=MUTED)

    doc.add_paragraph()
    meta = [
        ("Audience", "Thermal coal practitioner with limited Korea power-market background"),
        ("Model", "KPX Cost Stack Lab, static browser app"),
        ("App URL", "https://qrlow.github.io/kpx-cost-stack-lab/"),
        ("Repository", "https://github.com/qrlow/kpx-cost-stack-lab"),
        ("Date", "2026-06-25"),
    ]
    add_table(doc, ["Field", "Value"], meta, [1.4, 5.1])
    doc.add_page_break()


def add_toc(doc):
    doc.add_heading("How to Use This Course", level=1)
    add_para(doc, "This document explains the model as a course. It starts from institutional design, then builds toward the dispatch stack, then shows how coal market knowledge enters Korean power economics. You do not need prior Korea power-market knowledge, but the document assumes you are comfortable with thermal coal terms such as NAR, delivered price, freight, heat rate, and fuel switching.")
    add_callout(doc, "Core thesis", "Korea is not a simple merchant coal-versus-gas switching market. It is a cost-based, centrally scheduled power market where KPX builds a day-ahead schedule from assessed costs and system constraints, while government policy, nuclear availability, LNG procurement structure, air-quality rules, ETS treatment, and transmission/security constraints can all change the apparent coal/LNG economics.")
    doc.add_heading("Learning Objectives", level=2)
    add_bullets(doc, [
        "Explain why KPX dispatch is not equivalent to a free merchant bid stack.",
        "Translate coal and LNG commodity prices into KRW/kWh generation costs.",
        "Identify the difference between spot fuel economics and KPX-style recognized costs.",
        "Explain why nuclear availability, renewables, coal minimum load, and seasonal air-quality policy can dominate coal/LNG spreads.",
        "Use the KPX Cost Stack Lab controls without over-interpreting the output as confidential unit-level dispatch.",
        "Read model outputs as scenario diagnostics rather than market forecasts."
    ])
    doc.add_heading("Course Map", level=2)
    add_table(doc, ["Module", "Topic", "What you should be able to do"], [
        ("1", "Korea power-market architecture", "Name the core institutions and explain the single-buyer structure."),
        ("2", "KPX cost-based pool", "Explain GCAC, PSS, SMP, settlement, and why costs are assessed."),
        ("3", "KPX data baseline", "Read capacity, generation, SMP, and marginal fuel counts from KPX statistics."),
        ("4", "Fuel economics", "Convert coal and LNG prices into KRW/kWh using heat rates and FX."),
        ("5", "Residual load and constraints", "Explain why dispatch is about net load, not gross load alone."),
        ("6", "Policy overlays", "Explain fine-dust controls, K-ETS, carbon recognition, and 11th Basic Plan context."),
        ("7", "Model walkthrough", "Connect each app control to a market mechanism and equation."),
        ("8", "Limitations and upgrades", "Separate what the app demonstrates from what a production unit-commitment model would need."),
    ], [0.65, 2.1, 3.75])
    doc.add_page_break()


def add_module_1(doc):
    doc.add_heading("Module 1 - Korea Power-Market Architecture", level=1)
    add_para(doc, "Korea's electricity sector is best understood as a hybrid: generation has multiple participants, but transmission, distribution, retail supply, and the wholesale purchase role remain centered on KEPCO. The Korea Power Exchange (KPX) operates the wholesale market and the power system. KPX itself describes the market as a pool in which generators sell and KEPCO acts as the single buyer. [S2]")
    add_para(doc, "The IEA describes the structure as a day-ahead wholesale market run by KPX, with KEPCO controlling transmission, distribution, and retail supply. Generation comes from independent power producers and KEPCO subsidiaries. The six KEPCO generation subsidiaries remain central: Korea Hydro & Nuclear Power owns nuclear and hydro, while the other five gencos own much of the coal and gas fleet. [S7]")
    add_para(doc, "This institutional setup matters for coal switching because a generator is not simply deciding whether to bid coal at its own opportunity cost into a liberalized market. KPX schedules the system, KEPCO pays through the market settlement process, and generator costs are evaluated under market rules. That makes the Korean coal market a power-market-design problem as much as a fuel-price problem.")
    add_table(doc, ["Actor", "Role in the model", "Why it matters for coal/LNG switching"], [
        ("KPX", "Market operator, system operator, scheduler, settlement administrator", "Builds the day-ahead schedule and determines SMP under assessed-cost and constraint rules."),
        ("KEPCO", "Single buyer and retail supplier", "Purchases from the KPX pool and supplies end users; retail tariffs are regulated and do not freely float with fuel cost."),
        ("KEPCO gencos", "Dominant fleet owners", "Own much of the coal and LNG capacity; their units form much of the assessed cost stack."),
        ("KHNP", "Nuclear and hydro operator", "Nuclear availability strongly changes residual demand for coal and LNG."),
        ("IPPs/direct LNG importers", "Private generation and some self-use LNG import", "Can introduce plant-specific LNG economics that differ from the KOGAS average."),
        ("MOTIE/MCEE and policy bodies", "Planning, environmental, and ETS policy", "Coal caps, emissions costs, and the long-term supply plan shape the dispatchable fleet.")
    ], [1.35, 2.15, 3.0], "KPX market-process pages and IEA Korea Electricity Security Review. [S2], [S7]")
    add_callout(doc, "Coal-market intuition adjustment", "In a merchant market, a coal analyst asks whether coal variable cost is below gas variable cost. In Korea, ask whether the KPX-recognized cost block can move in the Price Setting Schedule after must-run supply, minimum stable coal output, nuclear availability, policy caps, LNG availability, and system constraints are applied.")


def add_module_2(doc):
    doc.add_heading("Module 2 - KPX Cost-Based Pool, GCAC, PSS, SMP, and Settlement", level=1)
    add_para(doc, "Korea's current wholesale market is a Cost-Based Pool. KPX states that both fixed costs, through capacity payment, and variable costs for each generating unit are examined monthly by the Generation Cost Assessment Committee (GCAC) based on documents submitted by generators. [S2]")
    add_para(doc, "The Cost Evaluation Committee matters because the supply curve used by KPX is not built from free price bids. KPX says that generation costs calculated from each generator's variable costs form the basis of market price in the Cost-Based Pool. The committee meets regularly and performs the cost evaluation function. [S5]")
    add_para(doc, "KPX determines market price one day ahead for each hour. The Price Setting Schedule (PSS) chooses generator operation and output levels to meet forecast demand while satisfying system constraints at the lowest generation cost. The highest-cost dispatched generator is designated as the marginal plant, and its System Marginal Price (SMP) is the market price for that time period. [S3]")
    add_para(doc, "KPX's trading system page adds an operational detail that is central to this model: the scheduler establishes a PSS, calculates SMP/BLMP a day ahead according to the trading day's demand forecast, and publishes an operation schedule considering fuel and transmission constraints. That is why a pure fuel spread is incomplete. [S4]")
    add_table(doc, ["Concept", "Plain-English definition", "How the app represents it"], [
        ("GCAC", "Committee process for assessed fixed and variable generator costs.", "The app uses representative recognized fuel costs and variable O&M rather than pure bid prices."),
        ("PSS", "Day-ahead cost-minimizing operating schedule subject to constraints.", "The app ranks available blocks by variable cost after must-run and policy constraints."),
        ("SMP", "System marginal price set by the marginal plant.", "The app reports the variable cost of the last dispatched block plus any reserve scarcity premium."),
        ("BLMP", "KPX also refers to BLMP in the trading-system scheduler context.", "The app does not model locational/network prices; it uses a single mainland-style stack."),
        ("Settlement", "KPX calculates trading amounts after trading from bid, metered, and dispatch data.", "The app focuses on dispatch logic, not cash settlement or uplift mechanics.")
    ], [1.0, 2.7, 2.8], "KPX market operation pages. [S2]-[S6]")
    doc.add_heading("Settlement Is Not the Same as Dispatch", level=2)
    add_para(doc, "KPX settlement is calculated after trading using bidding data, metered data, and dispatch data. KPX describes preliminary and final settlement steps, including notice to generators and KEPCO, review of objections, and payment transfer after KEPCO deposits funds. [S6]")
    add_para(doc, "For learning coal switching, do not confuse three layers: physical dispatch, price formation, and final cash settlement. The app models physical dispatch and marginal price formation. It does not model all settlement adjustments, capacity payments, constrained-on/off compensation, or generator-specific payment coefficients.")


def add_module_3(doc):
    doc.add_heading("Module 3 - The 2025 KPX Statistical Baseline", level=1)
    add_para(doc, "The app uses KPX's 2025 Power Market Statistics as its core baseline. KPX's explanatory notes define market-registered capacity as capacity registered in the power market at period end; KEPCO PPA and other non-market items are not necessarily included, so it is not identical to total national installed capacity. KPX defines trading volume as market-traded output net of station service and transformer losses, and again not necessarily total national generation. [S1]")
    add_table(doc, ["Fuel", "2025 market-registered capacity (MW)", "Capacity share", "2025 traded volume (GWh)", "Trading share"], [
        ("Nuclear", "26,050", "17.8%", "175,549", "32.1%"),
        ("Coal", "41,739", "28.6%", "164,790", "30.1%"),
        ("LNG", "48,388", "33.1%", "159,227", "29.1%"),
        ("Oil", "531", "0.4%", "160", "0.0%"),
        ("Pumped storage", "4,700", "3.2%", "4,414", "0.8%"),
        ("Renewable", "19,457", "13.3%", "39,697", "7.3%"),
        ("Other", "5,293", "3.6%", "3,325", "0.6%"),
        ("Total", "146,157", "100.0%", "547,163", "100.0%"),
    ], [1.2, 1.55, 1.05, 1.55, 1.15], "KPX 2025 Power Market Statistics. [S1]")
    add_para(doc, "Several patterns jump out. First, LNG had the largest market-registered capacity but not the largest energy output; nuclear and coal produced more energy relative to capacity because they typically run at higher utilization. Second, coal and LNG traded volumes were very close in 2025, which makes the coal/LNG switching question commercially important. Third, renewables were still a relatively small share of traded energy, but their hourly profile matters because solar changes the shape of residual demand.")
    add_table(doc, ["Metric", "2025 KPX value used in app", "Interpretation"], [
        ("Mainland weighted-average SMP", "112.73 KRW/kWh", "Anchor for checking whether modeled SMP is plausible, not a calibration target."),
        ("LNG marginal price-setting count", "7,292 of 8,741 intervals", "LNG set SMP in about 83.4% of counted intervals."),
        ("Coal marginal price-setting count", "1,413 of 8,741 intervals", "Coal set SMP in about 16.2% of counted intervals."),
        ("Hourly pattern", "Coal more frequent overnight; LNG more frequent most hours", "Supports the model design where LNG often marginal but coal still material."),
    ], [1.65, 1.7, 3.15], "KPX 2025 Power Market Statistics. [S1]")
    add_callout(doc, "Interpreting marginal counts", "A fuel's marginal price-setting count is not its generation share. Coal can generate a large amount without setting SMP if another resource is marginal. Conversely, LNG can set SMP frequently even if its generation share is similar to coal's.")


def add_module_4(doc):
    doc.add_heading("Module 4 - Translating Coal and LNG into KRW/kWh", level=1)
    add_para(doc, "A thermal coal analyst normally thinks in USD/t, kcal/kg, freight, sulphur, ash, and delivered basis. A power-market stack needs the same economics expressed as KRW/kWh. The app makes this translation explicitly so that coal knowledge is visible inside the dispatch model.")
    add_para(doc, "For coal, the conversion starts with calorific value. A tonne of 5,500 kcal/kg NAR coal contains about 23.03 GJ, or about 21.83 MMBtu. Therefore, USD/t divided by roughly 21.83 gives USD/MMBtu. Multiplying by USD/KRW gives KRW/MMBtu. Multiplying by heat rate in MMBtu/MWh gives KRW/MWh, and dividing by 1,000 gives KRW/kWh.")
    add_code(doc, "coal_USD_per_MMBtu = coal_USD_per_tonne / ((NAR_kcal_per_kg * 1000 * 4.1868 / 1,000,000) / 1.05505585262)")
    add_code(doc, "fuel_KRW_per_kWh = fuel_KRW_per_MMBtu * heat_rate_MMBtu_per_MWh / 1000")
    add_para(doc, "For LNG, market price is already quoted as USD/MMBtu in JKM-style terms, so the model uses USD/MMBtu multiplied by FX, then by plant heat rate. LNG combined-cycle plants have lower heat rates than coal units, but LNG fuel is usually more expensive per MMBtu. That is the fuel-switching tension.")
    add_table(doc, ["Input", "Coal example", "LNG example", "Why it matters"], [
        ("Commodity price", "98 USD/t at 5,500 kcal/kg NAR", "10.8 USD/MMBtu", "Sets fuel energy cost."),
        ("FX", "1,380 KRW/USD", "1,380 KRW/USD", "Korea imports both fuels; FX is a direct cost driver."),
        ("Heat rate", "8.75-9.55 MMBtu/MWh in model blocks", "6.55-7.75 MMBtu/MWh in model blocks", "Efficiency changes whether expensive LNG can compete with cheaper coal."),
        ("Variable O&M", "5.5-7.4 KRW/kWh", "3.8-5.8 KRW/kWh", "Non-fuel dispatch cost."),
        ("Carbon factor", "0.8384 tCO2e/MWh", "0.3800 tCO2e/MWh", "MOTIE average factors used in the updated app."),
    ], [1.25, 1.5, 1.45, 2.3], "Heat rates are representative model assumptions; emissions factors are from MOTIE 11th Basic Plan preview text. [S8]")
    doc.add_heading("Why Spot Price Is Not Enough", level=2)
    add_para(doc, "The app separates spot fuel prices from recognized fuel prices. Spot coal and spot LNG represent what a commodity analyst sees in the market. Recognized prices represent the lagged, documented, and evaluated fuel-cost basis that better approximates how a cost-based pool reorders generation. A spot coal selloff may not immediately push coal deeper into the KPX stack if procurement costs, inventory accounting, fuel-cost submission cycles, and GCAC evaluation lag the move.")
    add_para(doc, "This is especially important for Korea because most fuel is imported. Delivered cost includes not just commodity index price, but freight, quality adjustment, terminal and handling assumptions, taxes/levies, FX, contract terms, and accounting timing. The model therefore lets you test both an immediate spot world and a KPX-style recognized-cost world.")


def add_module_5(doc):
    doc.add_heading("Module 5 - LNG and Gas-Market Structure", level=1)
    add_para(doc, "LNG in Korea is not a single transparent spot fuel input. The IEA explains that Korea's gas industry is divided into wholesale and retail sectors. KOGAS manages the wholesale sector; regional city gas companies handle retail supply. KOGAS holds a dominant position across LNG imports, distribution, and the wholesale market, while large power and industrial consumers may directly import LNG for their own consumption if they meet requirements. [S13]")
    add_para(doc, "This matters for dispatch because a gas-fired plant's economics can differ by procurement route. A plant supplied under KOGAS tariffs, a private generator with direct LNG imports, and a CHP unit serving district heating can face different effective gas economics and operational obligations. The app simplifies this into three representative LNG blocks: high-efficiency CCGT, mid-merit CCGT, and CHP/peaking LNG.")
    add_table(doc, ["LNG channel", "Market feature", "Dispatch implication"], [
        ("KOGAS wholesale supply", "Dominant national LNG importer and wholesale supplier.", "Cost may reflect portfolio procurement, storage, and tariff timing rather than spot JKM alone."),
        ("Direct LNG import", "Large power/industrial users can import for own use but not resale.", "Some plants may have contract-specific fuel costs that differ from average KOGAS supply."),
        ("City gas / winter demand", "Residential and commercial heating creates seasonal gas demand.", "Winter gas demand can raise scarcity concerns and complicate power-sector switching."),
        ("CHP/district heating", "Some units serve heat load as well as electricity.", "Must-run heat obligations can keep gas units online even when pure power economics would differ."),
    ], [1.45, 2.25, 2.8], "IEA Natural Gas Supply Security in Korea. [S13]")
    add_para(doc, "The IEA also notes that power-generation gas demand is sensitive to nuclear expansion, renewable growth, hydrogen co-firing, and relative pricing. In a Korean switching model, LNG is both the main flexible competitor to coal and a security-of-supply exposure to global LNG markets. [S13]")


def add_module_6(doc):
    doc.add_heading("Module 6 - Residual Load, Must-Run Supply, and Physical Constraints", level=1)
    add_para(doc, "Dispatch is not simply a stack of every fuel from cheapest to most expensive. The schedule first has to serve net load: gross electricity demand minus non-dispatchable or must-take resources, plus operational requirements. Nuclear, renewables, CHP, pumped storage, reserve requirements, and grid constraints all shape the residual load that coal and LNG compete to serve.")
    add_para(doc, "Nuclear is central in Korea. The IEA's 2025 executive summary states that 26 reactors with 26 GW of installed capacity provide about a third of electricity, and nuclear remains one pillar of Korea's long-term energy and climate strategy. [S14] In the model, reducing nuclear availability raises residual demand and usually increases the need for LNG after coal reaches its modeled cap or availability limit.")
    add_para(doc, "Coal units have physical inflexibility. They have minimum stable output, slower start-up and ramping than gas turbines, minimum up/down constraints, and planned/forced outage schedules. The app approximates this with a coal minimum stable block plus flexible coal blocks. This is a simplification: a real unit-commitment model would track each unit's minimum output, ramp rate, start cost, and minimum up/down time.")
    add_para(doc, "Renewables change the shape rather than only the level of demand. Solar lowers midday net load and can increase evening ramp requirements. The IEA's Korea Electricity Security Review emphasizes that higher variable renewable energy raises flexibility requirements and changes net-load ramps. [S7]")
    add_table(doc, ["Constraint", "Real-world mechanism", "App approximation"], [
        ("Nuclear availability", "Planned/forced outages and utilization of low-variable-cost nuclear fleet.", "Nuclear availability percentage control."),
        ("Renewable weather", "Solar/wind output profile, curtailment, and weather variation.", "Hourly renewable profile scaled by renewable weather control."),
        ("Coal minimum load", "Coal units cannot be turned down to zero instantly.", "Mandatory coal minimum stable block."),
        ("Coal cap / fine dust", "Policy and air-quality restrictions can limit coal operation.", "Coal availability cap and fine-dust toggle."),
        ("Transmission constraints", "Local grid limits can force units on/off out of economic order.", "Transmission stress derates LNG availability and is discussed as a qualitative constraint."),
        ("Reserve requirement", "System reliability requires capacity beyond current load.", "Reserve target and scarcity premium.")
    ], [1.45, 2.55, 2.5])
    add_callout(doc, "Why LNG often sets SMP", "LNG often sits near the margin because it is flexible and generally higher variable cost than nuclear and coal. KPX 2025 statistics show LNG set SMP in the large majority of counted intervals. That does not mean LNG supplied most generation; it means LNG often supplied the marginal increment. [S1]")


def add_module_7(doc):
    doc.add_heading("Module 7 - Policy Overlays: Fine Dust, ETS, Carbon Recognition, and the 11th Basic Plan", level=1)
    doc.add_heading("Fine-Dust Seasonal Management", level=2)
    add_para(doc, "Korea's fine-dust policy can directly constrain coal operation. The Ministry of Climate, Energy and Environment explains that the Fine Dust Seasonal Management Plan implements stronger emission-reduction measures from December 1 to March 31, when fine-dust concentrations are typically higher. In the 5th plan, coal power measures included promoting shutdown of up to 15 units and output caps for up to 47 units at 80%. [S10]")
    add_para(doc, "This is not a fuel-price switch. Coal can be backed down because air-quality policy says so, even if coal is cheaper than LNG. The app represents this with a fine-dust toggle and a coal availability cap.")
    doc.add_heading("K-ETS and Carbon Recognition", level=2)
    add_para(doc, "The Korea Emissions Trading System launched in 2015 and covers power, industry, buildings, waste, transport, domestic aviation, and maritime sectors. ICAP reports that K-ETS covered 813 large emitters for the 2025 compliance year and that Phase 4 runs from 2026 to 2030. For power generation, allowance auctioning is scheduled to rise from 15% in 2026 to 50% by 2030. [S11]")
    add_para(doc, "Carbon pricing does not automatically behave the same way in every electricity market. The IEA's discussion of ETS in Korea's power sector highlights that Korea's cost-based wholesale market and regulated retail prices historically limited how carbon costs influenced dispatch and retail price signals. It describes the policy challenge of reflecting ETS carbon prices in the power generation plan through environmental dispatch or similar mechanisms. [S12]")
    add_para(doc, "The updated model therefore separates carbon price from carbon recognition. Carbon price is the allowance value in KRW/tCO2e. Carbon recognition is the percentage of that price reflected in the modeled dispatch variable cost. A 100% recognition setting asks what happens when the full carbon value is in the stack; a 0% setting asks what the fuel stack looks like when carbon is not dispatch-relevant.")
    add_code(doc, "carbon_cost_KRW_per_kWh = emissions_tCO2e_per_MWh * carbon_KRW_per_tCO2e * carbon_recognition_pct / 100 / 1000")
    add_para(doc, "The model now uses MOTIE's Korea-specific power-sector average 2020-2022 emissions factors: coal 0.8384 tCO2e/MWh and LNG 0.3800 tCO2e/MWh. These are more appropriate than generic coal and gas factors for a Korea-focused teaching model. [S8]")
    doc.add_heading("11th Basic Plan", level=2)
    add_para(doc, "The 11th Basic Plan for Long-term Electricity Supply and Demand covers 2024-2038. KDI's policy release states that the plan responds to future electricity demand from advanced industries, data centers, and electrification while considering supply stability, efficiency, and carbon neutrality. It also confirms two new large nuclear units, one SMR, and efforts to add renewable capacity at an average of 7 GW per year through 2030. [S9]")
    add_para(doc, "The plan matters for the model because it changes the long-run fleet and therefore the future cost stack. More nuclear and renewables reduce residual fossil generation in many hours; more storage and flexible resources change the ramping problem; coal retirements and conversions reduce coal's structural role; LNG can remain important as flexibility even if its annual energy declines.")
    add_table(doc, ["Policy lever", "Directional effect on coal burn", "Directional effect on LNG burn", "Model control"], [
        ("Fine-dust coal cap", "Lower coal output regardless of fuel spread", "Often higher LNG if load must still be met", "Fine-dust toggle; coal cap"),
        ("Carbon recognition", "Raises coal cost more than LNG because coal has higher CO2 intensity", "Can improve LNG relative position", "Carbon cost; carbon recognition"),
        ("Nuclear availability", "High nuclear suppresses fossil residual demand", "Low nuclear raises fossil residual demand", "Nuclear availability"),
        ("Renewables", "Midday solar can displace coal if coal can turn down", "Raises need for flexible ramping in some hours", "Renewable weather; hour"),
        ("Coal retirements/conversions", "Lower structural coal capacity", "May raise LNG flexibility role unless alternatives scale", "Coal availability cap"),
    ], [1.55, 1.7, 1.7, 1.55])


def add_module_8(doc):
    doc.add_heading("Module 8 - Walking Through the KPX Cost Stack Lab", level=1)
    add_para(doc, "The app is a one-hour dispatch simulator. It is not a forecast engine. It asks a controlled question: given this load, fuel-cost assumption, carbon treatment, availability, and policy state, which representative generation blocks would serve demand and which block would be marginal?")
    add_table(doc, ["Control", "Meaning", "Common interpretation mistake"], [
        ("Hour", "Selects hourly SMP comparison and renewable profile.", "Assuming all hours have the same renewable and marginal-fuel behavior."),
        ("Load", "Gross system load for the modeled hour.", "Ignoring that residual load after nuclear/renewables is what coal and LNG compete for."),
        ("Recognized Coal / LNG", "KPX-style lagged fuel-cost input.", "Treating it as a live market quote."),
        ("Spot Coal / LNG", "Immediate market-price counterfactual.", "Assuming spot instantly changes dispatch under a cost-based pool."),
        ("USD/KRW", "FX conversion for imported fuels.", "Forgetting that Korea's imported fuel cost is highly FX-sensitive."),
        ("Carbon Cost", "Allowance price or shadow carbon value.", "Assuming it is always fully reflected in dispatch."),
        ("Carbon Recognition", "Percentage of carbon value included in variable cost.", "Ignoring Korea's regulated market-design and pass-through issues."),
        ("Coal Availability Cap", "Policy/availability derate for coal blocks.", "Reading it as a precise unit-level outage schedule."),
        ("Nuclear Availability", "Available share of nuclear capacity.", "Underestimating nuclear as a driver of fossil residual demand."),
        ("Renewable Weather", "Scales the hourly renewable output profile.", "Confusing installed renewable capacity with hourly output."),
    ], [1.45, 2.35, 2.7])
    doc.add_heading("Default Scenario After the Audit Update", level=2)
    add_table(doc, ["Output", "Default value", "Interpretation"], [
        ("Modeled SMP", "107.8 KRW/kWh", "Marginal block variable cost plus scarcity premium under current defaults."),
        ("Marginal block", "High-efficiency LNG CCGT", "LNG sets price in the default hour despite coal producing more energy."),
        ("Coal dispatch", "32,556 MW", "44.6% of modeled dispatched supply."),
        ("LNG dispatch", "12,225 MW", "16.8% of modeled dispatched supply."),
        ("Nuclear dispatch", "22,403 MW", "Large low-variable-cost base supply reduces fossil residual load."),
        ("Policy counterfactual", "Coal cap holds coal about 9,183 MW below unconstrained coal dispatch", "The default coal cap is binding and pushes equivalent MW toward LNG."),
    ], [1.45, 1.65, 3.4], "Computed from the updated model on 2026-06-25. [S19]")
    add_para(doc, "The key diagnostic is not whether the modeled SMP exactly equals KPX's latest hourly value. The key diagnostic is which causal lever moves coal and LNG. If toggling spot prices does little but relaxing the coal cap moves thousands of MW, then the scenario is policy-constrained rather than price-constrained.")
    doc.add_heading("How to Read the Switching Diagnostics", level=2)
    add_bullets(doc, [
        "Coal policy/cap effect: compares the base case to an unconstrained coal-cap case. Negative coal delta means the cap suppresses coal output.",
        "Spot-vs-recognized fuel effect: compares KPX-style lagged fuel inputs to immediate spot fuel inputs. A small effect means the spot spread does not reorder the representative stack under the chosen scenario.",
        "Nuclear outage sensitivity: lowers nuclear availability and shows how residual load is reallocated. If coal is capped or fully used, the extra residual load tends to go to LNG.",
    ])


def add_module_9(doc):
    doc.add_heading("Module 9 - What the Model Does Not Do", level=1)
    add_para(doc, "A useful teaching model must be honest about its boundary. The KPX Cost Stack Lab is a representative cost-stack simulator. It is not a confidential KPX production model, a full unit-commitment optimizer, a network model, or a settlement engine.")
    add_table(doc, ["Missing production feature", "Why it matters", "How to upgrade"], [
        ("Unit-level fleet data", "Actual units have specific heat rates, outage schedules, min loads, ramp rates, and fuel contracts.", "Build a plant/unit table from public capacity lists and outage data where available."),
        ("Commitment optimization", "Start costs, minimum up/down time, and ramp rates can keep a unit online even when its marginal cost is high.", "Use a mixed-integer unit-commitment library or power-system optimizer."),
        ("Network constraints", "Transmission can force local generation out of merit order.", "Add zones or nodal approximations; incorporate constraint shadow prices if available."),
        ("Full settlement logic", "Payments include more than SMP times energy.", "Model capacity payments, make-whole payments, constrained-on/off compensation, and adjustment coefficients if public data permit."),
        ("Live data ingestion", "KPX live pages and commodity prices change daily.", "Add automated data fetches, versioned snapshots, and validation logs."),
        ("Fuel-quality economics", "Coal value depends on CV, sulphur, ash, moisture, freight, and plant constraints.", "Add a coal cargo calculator feeding delivered KRW/MMBtu by plant type."),
    ], [1.65, 2.3, 2.55])
    add_callout(doc, "Proper use", "Use the app to teach mechanisms and test scenarios. Do not use it to price real cargoes, forecast KPX dispatch, value a plant, or make financial trading decisions without replacing representative assumptions with verified unit-level and contract-level data.")


def add_model_audit(doc):
    doc.add_heading("Model Audit and Updates Made During This Work", level=1)
    add_para(doc, "While preparing this course, the model was audited against the referenced sources. Two improvements were made.")
    add_numbers(doc, [
        "Emissions factors were changed from generic coal/LNG factors to MOTIE 11th Basic Plan Korea-specific average factors: coal 0.8384 tCO2e/MWh and LNG 0.3800 tCO2e/MWh. [S8]",
        "A carbon-recognition percentage control was added so the user can distinguish carbon allowance price from the share reflected in dispatch variable cost. This aligns the app with IEA's warning that carbon-cost reflection in Korea's regulated cost-based electricity market is a market-design question, not automatic pass-through. [S12]",
        "The default carbon-price input was updated to KRW 16,700/tCO2e using a public OPIS assessment reported by MarketWatch/Dow Jones. It remains an illustrative scenario input, not a live price feed. [S16]",
        "Static asset query versions were bumped so GitHub Pages visitors do not keep the old cached model files."
    ])
    add_para(doc, "No change was made to the app's fundamental scope. It remains a representative teaching model with transparent assumptions. The correct next improvement would be adding unit-level plant blocks and live KPX/commodity data ingestion with a versioned data-quality log.")


def add_appendix(doc):
    doc.add_heading("Appendix A - Equations and Units", level=1)
    add_table(doc, ["Equation", "Use"], [
        ("GJ/t = NAR kcal/kg * 1000 kg/t * 4.1868 kJ/kcal / 1,000,000", "Convert coal calorific value to energy per tonne."),
        ("MMBtu/t = GJ/t / 1.05505585262", "Convert coal energy to MMBtu."),
        ("Coal USD/MMBtu = USD/t / MMBtu/t", "Compare coal to gas on energy basis."),
        ("Fuel KRW/kWh = USD/MMBtu * FX * heat rate / 1000", "Translate imported fuel cost into generation cost."),
        ("Carbon KRW/kWh = emissions tCO2e/MWh * carbon KRW/t * recognition pct / 100 / 1000", "Translate ETS/shadow carbon into variable cost."),
        ("Variable cost = fuel cost + variable O&M + recognized carbon cost", "Block-level dispatch ranking."),
        ("Reserve margin = (available capacity - load) / load", "Approximate adequacy pressure in the model."),
    ], [3.6, 2.9])
    doc.add_heading("Appendix B - Data Dictionary for the Static App", level=1)
    add_table(doc, ["Data item", "Meaning", "Source"], [
        ("marketRegisteredCapacityMw", "KPX market-registered capacity by fuel in MW.", "KPX 2025 statistics [S1]"),
        ("tradingVolumeGwh", "KPX market-traded generation by fuel in GWh.", "KPX 2025 statistics [S1]"),
        ("mainlandSmpByHourKrwPerKwh", "2025 mainland hourly weighted-average SMP profile.", "KPX 2025 statistics [S1]"),
        ("marginalPriceSetCountsByFuel", "Counts of intervals in which each fuel set SMP.", "KPX 2025 statistics [S1]"),
        ("latestKpxSnapshot", "Static snapshot of KPX SMP and supply-demand pages used when the app was built.", "KPX live pages [S17], [S18]"),
        ("representativeBlocks", "Teaching approximation of Korean generation fleet blocks.", "Calibrated to KPX capacity/energy structure; not confidential unit data."),
        ("emissionsTonnePerMwh", "CO2e factors by fuel used for carbon cost.", "MOTIE for coal/LNG; representative for oil/other [S8]"),
    ], [1.7, 2.7, 2.1])
    doc.add_heading("Appendix C - Suggested Exercises", level=1)
    add_numbers(doc, [
        "Set carbon recognition to 0%. Record coal dispatch, LNG dispatch, and modeled SMP. Then set it to 100%. Explain which fuel moves and why.",
        "Reduce nuclear availability by 15 percentage points. If coal does not increase much, identify which constraint prevents coal from absorbing residual load.",
        "Turn on fine-dust season and lower the coal cap. Explain why LNG may run even when coal remains cheaper on fuel cost.",
        "Switch from recognized to spot fuel mode. If dispatch barely changes, explain why the representative stack order did not cross.",
        "Raise LNG price and lower coal price. Then relax coal cap to 100%. Explain the difference between price-driven and policy-driven switching.",
        "Write a short analyst note decomposing a change in coal burn into demand, nuclear, renewables, fuel spread, carbon, and policy-cap effects."
    ])


def add_references(doc):
    doc.add_heading("References", level=1)
    for sid, title, url in SOURCES:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"[{sid}] {title} ")
        set_run(r, bold=True)
        r2 = p.add_run(url)
        set_run(r2, color=RGBColor(0, 102, 204))


def build():
    doc = Document()
    configure_styles(doc)
    set_header_footer(doc)
    add_cover(doc)
    add_toc(doc)
    add_module_1(doc)
    add_module_2(doc)
    add_module_3(doc)
    add_module_4(doc)
    add_module_5(doc)
    add_module_6(doc)
    add_module_7(doc)
    add_module_8(doc)
    add_module_9(doc)
    add_model_audit(doc)
    add_appendix(doc)
    add_references(doc)
    doc.core_properties.title = "Understanding Korea's KPX Cost Stack and Thermal Coal-to-LNG Switching"
    doc.core_properties.subject = "Korea power-market course and KPX Cost Stack Lab guide"
    doc.core_properties.author = "OpenAI Codex"
    doc.core_properties.comments = "Generated with compact_reference_guide preset; sources listed in document."
    doc.save(OUT)


if __name__ == "__main__":
    build()
